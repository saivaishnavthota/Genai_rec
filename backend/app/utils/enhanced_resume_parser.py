"""
Enhanced Resume Parser with improved NLP-based extraction
This module provides better resume parsing using advanced pattern matching and context analysis
"""

import re
import PyPDF2
from docx import Document
from typing import Dict, List, Any, Tuple
import logging
from collections import Counter

logger = logging.getLogger(__name__)

# Expanded skill patterns with categories
SKILL_CATEGORIES = {
    'programming_languages': [
        r'\b(?:Python|Java|JavaScript|TypeScript|C\+\+|C#|PHP|Ruby|Go|Rust|Swift|Kotlin|Scala|R|MATLAB|Perl|Shell|Bash)\b',
    ],
    'web_frameworks': [
        r'\b(?:React|Angular|Vue\.?js|Svelte|Next\.?js|Nuxt\.?js|Gatsby)\b',
        r'\b(?:Node\.?js|Express\.?js|Django|Flask|FastAPI|Spring|Spring Boot|Laravel|Rails|ASP\.NET)\b',
    ],
    'databases': [
        r'\b(?:MySQL|PostgreSQL|MongoDB|Redis|Elasticsearch|Cassandra|DynamoDB|Oracle|SQL Server|MariaDB|SQLite)\b',
        r'\b(?:Neo4j|CouchDB|Firebase|Firestore|Supabase)\b',
    ],
    'cloud_devops': [
        r'\b(?:AWS|Azure|GCP|Google Cloud|Heroku|DigitalOcean|Linode)\b',
        r'\b(?:Docker|Kubernetes|K8s|Jenkins|GitLab CI|GitHub Actions|CircleCI|Travis CI)\b',
        r'\b(?:Terraform|Ansible|Chef|Puppet|CloudFormation)\b',
    ],
    'frontend': [
        r'\b(?:HTML5?|CSS3?|SCSS|Sass|Less|Bootstrap|Tailwind|Material-UI|Ant Design|Chakra UI)\b',
        r'\b(?:Webpack|Vite|Rollup|Parcel|Babel|ESLint|Prettier)\b',
    ],
    'data_science': [
        r'\b(?:Machine Learning|Deep Learning|AI|Artificial Intelligence|NLP|Computer Vision|Neural Networks)\b',
        r'\b(?:TensorFlow|PyTorch|Keras|Scikit-learn|Pandas|NumPy|SciPy|Matplotlib|Seaborn|Plotly)\b',
        r'\b(?:Data Analysis|Data Science|Statistics|Statistical Analysis|Predictive Modeling)\b',
    ],
    'mobile': [
        r'\b(?:React Native|Flutter|Ionic|Xamarin|SwiftUI|Jetpack Compose)\b',
        r'\b(?:iOS Development|Android Development|Mobile Development)\b',
    ],
    'tools': [
        r'\b(?:Git|GitHub|GitLab|Bitbucket|SVN|Mercurial)\b',
        r'\b(?:Jira|Confluence|Trello|Asana|Monday\.com)\b',
        r'\b(?:VS Code|IntelliJ|PyCharm|Eclipse|Visual Studio)\b',
    ],
    'methodologies': [
        r'\b(?:Agile|Scrum|Kanban|Waterfall|DevOps|CI/CD|TDD|BDD|Microservices|REST|GraphQL|gRPC)\b',
    ],
    'soft_skills': [
        r'\b(?:Leadership|Team Management|Communication|Problem Solving|Critical Thinking|Collaboration)\b',
        r'\b(?:Project Management|Time Management|Analytical Skills|Creativity|Adaptability)\b',
    ]
}

# Enhanced job title patterns with seniority levels
JOB_TITLE_PATTERNS = [
    # Software Engineering
    r'(?:Senior|Lead|Principal|Staff|Junior|Associate)?\s*(?:Software|Full[- ]?Stack|Backend|Frontend|Web)\s+(?:Engineer|Developer|Programmer)',
    r'(?:Senior|Lead|Principal|Staff|Junior)?\s*(?:DevOps|Site Reliability|Platform|Cloud)\s+Engineer',
    r'(?:Senior|Lead|Principal)?\s*(?:Solutions?|Enterprise|Technical)\s+Architect',
    
    # Data & AI
    r'(?:Senior|Lead|Principal|Junior)?\s*(?:Data|ML|Machine Learning|AI)\s+(?:Scientist|Engineer|Analyst)',
    r'(?:Senior|Lead)?\s*(?:Data|Business|Financial)\s+Analyst',
    
    # Management
    r'(?:Engineering|Product|Project|Program|Technical)\s+(?:Manager|Lead|Director)',
    r'(?:CTO|VP|Head)\s+of\s+(?:Engineering|Technology|Product)',
    
    # Design & Product
    r'(?:Senior|Lead|Junior)?\s*(?:UI/UX|Product|Graphic|Visual)\s+Designer',
    r'(?:Senior|Associate)?\s*Product\s+(?:Manager|Owner)',
    
    # QA & Testing
    r'(?:Senior|Lead|Junior)?\s*(?:QA|Quality Assurance|Test)\s+(?:Engineer|Analyst|Automation Engineer)',
    
    # Security
    r'(?:Senior|Lead)?\s*(?:Security|Cybersecurity|Information Security)\s+(?:Engineer|Analyst|Architect)',
]

# Enhanced education patterns
EDUCATION_PATTERNS = {
    'degrees': [
        r'\b(?:Bachelor(?:\'s)?|B\.?S\.?|B\.?A\.?|B\.?Tech|B\.?E\.?|B\.?Sc\.?)\b',
        r'\b(?:Master(?:\'s)?|M\.?S\.?|M\.?A\.?|M\.?Tech|M\.?E\.?|M\.?Sc\.?|MBA)\b',
        r'\b(?:Ph\.?D\.?|Doctorate|Doctoral)\b',
        r'\b(?:Associate|A\.?S\.?|A\.?A\.?)\b',
    ],
    'fields': [
        r'\b(?:Computer Science|CS|Software Engineering|Information Technology|IT)\b',
        r'\b(?:Electrical|Electronics|Mechanical|Civil|Chemical)\s+Engineering\b',
        r'\b(?:Mathematics|Statistics|Physics|Chemistry|Biology)\b',
        r'\b(?:Business Administration|Management|Finance|Economics|Accounting)\b',
        r'\b(?:Data Science|Artificial Intelligence|Machine Learning)\b',
    ]
}

# Enhanced certification patterns
CERTIFICATION_PATTERNS = [
    # Cloud
    r'\b(?:AWS|Amazon Web Services)\s+Certified\s+(?:Solutions Architect|Developer|SysOps|DevOps|Security|Data Analytics|Machine Learning)',
    r'\b(?:Microsoft|Azure)\s+Certified\s+(?:Solutions Architect|Developer|Administrator|DevOps|Data Engineer|AI Engineer)',
    r'\b(?:Google Cloud|GCP)\s+Certified\s+(?:Professional|Associate)\s+(?:Cloud Architect|Data Engineer|DevOps Engineer)',
    
    # Project Management
    r'\b(?:PMP|Project Management Professional|PRINCE2|Certified ScrumMaster|CSM|Professional Scrum Master|PSM)\b',
    r'\b(?:Agile Certified Practitioner|PMI-ACP|SAFe|Scaled Agile)\b',
    
    # Security
    r'\b(?:CISSP|CISM|CISA|CEH|CompTIA Security\+|OSCP)\b',
    
    # IT & Networking
    r'\b(?:CCNA|CCNP|CCIE|CompTIA A\+|CompTIA Network\+|ITIL)\b',
    
    # Data & Analytics
    r'\b(?:Certified Analytics Professional|CAP|Cloudera|Databricks|Tableau|Power BI)\b',
]


def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF file with improved handling"""
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        return text.strip()
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {e}")
        return ""


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX file with improved handling"""
    try:
        doc = Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + " "
                text += "\n"
        
        return text.strip()
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {e}")
        return ""


def extract_skills_enhanced(text: str) -> Tuple[List[str], Dict[str, List[str]]]:
    """
    Extract skills with categorization
    Returns: (all_skills, categorized_skills)
    """
    all_skills = []
    categorized_skills = {}
    
    for category, patterns in SKILL_CATEGORIES.items():
        category_skills = []
        for pattern in patterns:
            matches = re.findall(pattern, text, re.IGNORECASE)
            category_skills.extend(matches)
        
        # Remove duplicates while preserving case
        unique_skills = list(dict.fromkeys(category_skills))
        if unique_skills:
            categorized_skills[category] = unique_skills
            all_skills.extend(unique_skills)
    
    # Remove duplicates from all_skills
    all_skills = list(dict.fromkeys(all_skills))
    
    return all_skills, categorized_skills


def extract_experience_enhanced(text: str) -> List[Dict[str, Any]]:
    """Extract work experience with improved parsing"""
    experiences = []
    lines = text.split('\n')
    
    # Find experience section
    exp_section_start = -1
    for i, line in enumerate(lines):
        if re.search(r'\b(?:experience|work history|employment|professional experience)\b', line, re.IGNORECASE):
            exp_section_start = i
            break
    
    if exp_section_start == -1:
        # No explicit section, try to find job titles throughout
        exp_section_start = 0
    
    # Extract experiences
    i = exp_section_start
    while i < len(lines):
        line = lines[i].strip()
        
        # Check if line matches job title pattern
        for pattern in JOB_TITLE_PATTERNS:
            if re.search(pattern, line, re.IGNORECASE):
                exp_entry = {
                    'title': line,
                    'company': '',
                    'duration': '',
                    'location': '',
                    'description': []
                }
                
                # Look ahead for company, dates, location
                for j in range(i + 1, min(i + 5, len(lines))):
                    next_line = lines[j].strip()
                    
                    if not next_line:
                        continue
                    
                    # Check for date patterns
                    date_match = re.search(
                        r'(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{4}\s*[-–—]\s*(?:(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)[a-z]*\.?\s+\d{4}|Present|Current)',
                        next_line,
                        re.IGNORECASE
                    )
                    if date_match:
                        exp_entry['duration'] = date_match.group(0)
                        continue
                    
                    # Check for location
                    location_match = re.search(r'\b[A-Z][a-z]+(?:\s+[A-Z][a-z]+)*,\s*[A-Z]{2}\b', next_line)
                    if location_match:
                        exp_entry['location'] = location_match.group(0)
                        continue
                    
                    # If not date or location, likely company name
                    if not exp_entry['company'] and not re.search(r'^[•\-\*]', next_line):
                        exp_entry['company'] = next_line
                        continue
                    
                    # Collect bullet points as description
                    if re.search(r'^[•\-\*]', next_line):
                        exp_entry['description'].append(next_line.lstrip('•-* '))
                
                experiences.append(exp_entry)
                break
        
        i += 1
    
    return experiences


def extract_education_enhanced(text: str) -> List[Dict[str, Any]]:
    """Extract education with improved parsing"""
    education = []
    lines = text.split('\n')
    
    # Find education section
    edu_section_start = -1
    for i, line in enumerate(lines):
        if re.search(r'\b(?:education|academic|qualifications)\b', line, re.IGNORECASE):
            edu_section_start = i
            break
    
    if edu_section_start == -1:
        edu_section_start = 0
    
    # Extract education entries
    for i in range(edu_section_start, len(lines)):
        line = lines[i].strip()
        
        # Check for degree patterns
        for degree_pattern in EDUCATION_PATTERNS['degrees']:
            if re.search(degree_pattern, line, re.IGNORECASE):
                edu_entry = {
                    'degree': line,
                    'field': '',
                    'institution': '',
                    'year': '',
                    'gpa': ''
                }
                
                # Extract field of study
                for field_pattern in EDUCATION_PATTERNS['fields']:
                    field_match = re.search(field_pattern, line, re.IGNORECASE)
                    if field_match:
                        edu_entry['field'] = field_match.group(0)
                        break
                
                # Look ahead for institution and year
                for j in range(i + 1, min(i + 3, len(lines))):
                    next_line = lines[j].strip()
                    
                    # Extract year
                    year_match = re.search(r'\b(19|20)\d{2}\b', next_line)
                    if year_match and not edu_entry['year']:
                        edu_entry['year'] = year_match.group(0)
                    
                    # Extract GPA
                    gpa_match = re.search(r'GPA:?\s*(\d+\.?\d*)\s*/\s*(\d+\.?\d*)', next_line, re.IGNORECASE)
                    if gpa_match:
                        edu_entry['gpa'] = f"{gpa_match.group(1)}/{gpa_match.group(2)}"
                    
                    # Institution (if not already found)
                    if not edu_entry['institution'] and not year_match and not gpa_match:
                        edu_entry['institution'] = next_line
                
                education.append(edu_entry)
                break
    
    return education


def extract_certifications_enhanced(text: str) -> List[str]:
    """Extract certifications with improved patterns"""
    certifications = []
    
    for pattern in CERTIFICATION_PATTERNS:
        matches = re.findall(pattern, text, re.IGNORECASE)
        certifications.extend(matches)
    
    # Remove duplicates
    certifications = list(dict.fromkeys(certifications))
    
    return certifications


def calculate_ats_score_enhanced(text: str, filename: str, parsed_data: Dict[str, Any]) -> Dict[str, float]:
    """
    Calculate detailed ATS compatibility score based on industry standards
    Returns breakdown of scores
    """
    scores = {
        'format_score': 0.0,
        'structure_score': 0.0,
        'keyword_score': 0.0,
        'readability_score': 0.0
    }
    
    # 1. Format Score (25 points)
    if filename.lower().endswith('.pdf'):
        scores['format_score'] = 25.0
    elif filename.lower().endswith('.docx'):
        scores['format_score'] = 25.0
    elif filename.lower().endswith('.doc'):
        scores['format_score'] = 15.0
    else:
        scores['format_score'] = 5.0
    
    # 2. Structure Score (25 points) - Check for standard sections
    structure_points = 0
    required_sections = {
        'contact': r'\b(?:email|phone|linkedin|github)\b',
        'experience': r'\b(?:experience|work history|employment)\b',
        'education': r'\b(?:education|academic|degree)\b',
        'skills': r'\b(?:skills|technical skills|competencies)\b',
        'summary': r'\b(?:summary|objective|profile|about)\b'
    }
    
    for section, pattern in required_sections.items():
        if re.search(pattern, text, re.IGNORECASE):
            structure_points += 5
    
    scores['structure_score'] = structure_points
    
    # 3. Keyword Score (25 points) - Based on extracted skills and content richness
    skills_count = len(parsed_data.get('parsed_skills', []))
    if skills_count >= 15:
        scores['keyword_score'] = 25.0
    elif skills_count >= 10:
        scores['keyword_score'] = 20.0
    elif skills_count >= 5:
        scores['keyword_score'] = 15.0
    else:
        scores['keyword_score'] = 10.0
    
    # 4. Readability Score (25 points) - Check for good formatting indicators
    readability_points = 0
    
    # Has bullet points
    if re.search(r'[•\-\*]', text):
        readability_points += 5
    
    # Has proper dates
    if re.search(r'\b(?:19|20)\d{2}\b', text):
        readability_points += 5
    
    # Has contact info
    if re.search(r'[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}', text):
        readability_points += 5
    
    # Has phone number
    if re.search(r'\b\d{3}[-.]?\d{3}[-.]?\d{4}\b', text):
        readability_points += 5
    
    # Not too short or too long
    word_count = len(text.split())
    if 300 <= word_count <= 1500:
        readability_points += 5
    
    scores['readability_score'] = readability_points
    
    return scores


def parse_resume_enhanced(file_path: str, filename: str) -> Dict[str, Any]:
    """
    Enhanced resume parser with better extraction and categorization
    """
    try:
        # Extract text based on file type
        if filename.lower().endswith('.pdf'):
            text = extract_text_from_pdf(file_path)
        elif filename.lower().endswith(('.docx', '.doc')):
            text = extract_text_from_docx(file_path)
        else:
            logger.warning(f"Unsupported file format: {filename}")
            return _empty_parse_result()
        
        if not text or len(text.strip()) < 50:
            logger.warning(f"Insufficient text extracted from {filename}")
            return _empty_parse_result()
        
        # Extract information with enhanced methods
        all_skills, categorized_skills = extract_skills_enhanced(text)
        experience = extract_experience_enhanced(text)
        education = extract_education_enhanced(text)
        certifications = extract_certifications_enhanced(text)
        
        # Prepare parsed data for ATS scoring
        parsed_data = {
            'parsed_skills': all_skills,
            'parsed_experience': experience,
            'parsed_education': education,
            'parsed_certifications': certifications
        }
        
        # Calculate detailed ATS scores
        ats_scores = calculate_ats_score_enhanced(text, filename, parsed_data)
        overall_ats_score = sum(ats_scores.values())
        
        return {
            'parsed_skills': all_skills,
            'categorized_skills': categorized_skills,
            'parsed_experience': experience,
            'parsed_education': education,
            'parsed_certifications': certifications,
            'ats_score': overall_ats_score,
            'ats_score_breakdown': ats_scores,
            'raw_text': text[:2000],  # First 2000 chars for reference
            'word_count': len(text.split()),
            'has_contact_info': bool(re.search(r'[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}', text))
        }
        
    except Exception as e:
        logger.error(f"Error parsing resume {filename}: {e}", exc_info=True)
        return _empty_parse_result()


def _empty_parse_result() -> Dict[str, Any]:
    """Return empty parse result structure"""
    return {
        'parsed_skills': [],
        'categorized_skills': {},
        'parsed_experience': [],
        'parsed_education': [],
        'parsed_certifications': [],
        'ats_score': 0.0,
        'ats_score_breakdown': {
            'format_score': 0.0,
            'structure_score': 0.0,
            'keyword_score': 0.0,
            'readability_score': 0.0
        },
        'raw_text': '',
        'word_count': 0,
        'has_contact_info': False
    }
